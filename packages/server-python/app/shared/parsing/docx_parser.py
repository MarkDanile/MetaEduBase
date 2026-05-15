"""DOCX text and heading extraction using python-docx."""

from __future__ import annotations

from app.shared.parsing.pdf_parser import ParsedDocument, DocumentSection


def extract_docx_text(file_path: str) -> ParsedDocument:
    """Extract structured text from a DOCX file."""
    from docx import Document

    doc = Document(file_path)
    sections: list[DocumentSection] = []
    full_text_parts: list[str] = []

    current_title = ""
    current_level = 0
    current_content_parts: list[str] = []
    section_counter: dict[int, int] = {}

    _HEADING_STYLE = {
        "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
        "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
        "标题 1": 1, "标题 2": 2, "标题 3": 3,
        "标题 4": 4, "标题 5": 5, "标题 6": 6,
    }

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        heading_level = _HEADING_STYLE.get(style_name, 0)
        text = para.text.strip()

        if not text:
            continue

        if heading_level > 0:
            if current_title:
                content = "\n".join(current_content_parts).strip()
                path = _build_section_path(section_counter, heading_level)
                sections.append(DocumentSection(
                    title=current_title,
                    level=current_level,
                    content=content,
                    page=0,
                    path=path,
                ))
                full_text_parts.append(f"## {current_title}\n{content}")

            current_title = text
            current_level = heading_level
            current_content_parts = []
            section_counter[heading_level] = section_counter.get(heading_level, 0) + 1
        else:
            current_content_parts.append(text)

    if current_title:
        content = "\n".join(current_content_parts).strip()
        path = _build_section_path(section_counter, current_level)
        sections.append(DocumentSection(
            title=current_title,
            level=current_level,
            content=content,
            page=0,
            path=path,
        ))
        full_text_parts.append(f"## {current_title}\n{content}")

    if not sections:
        all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if all_text.strip():
            sections.append(DocumentSection(title="", level=0, content=all_text.strip(), page=0, path=""))
            full_text_parts.append(all_text.strip())

    return ParsedDocument(sections=sections, full_text="\n\n".join(full_text_parts))


def _build_section_path(counter: dict[int, int], level: int) -> str:
    parts = []
    for lvl in sorted(counter.keys()):
        if lvl <= level:
            parts.append(str(counter[lvl]))
    return ".".join(parts) if parts else ""
